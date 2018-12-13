import os
import docx
import json
import time

from django.http import JsonResponse, HttpResponse

from QuestionBank.settings import QB_MEDIA_DIR
from QuestionBank.models import User, UserProfile, Subject, Choice, Fill, Judge, Discuss

PAPER_DIR = os.path.join(QB_MEDIA_DIR, 'paper')

DOCX_CONTENT_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


def create_paper(request):
    user = User.objects.get(username=request.POST['openid'])
    profile = UserProfile.objects.get(user=user)

    # 检测是否教师
    if not profile.isTeacher:
        response = {'status': 'fail', 'errMsg': 'permission denied.'}
        return JsonResponse(response)

    doc = docx.Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')

    # 获取学科信息
    subject = Subject.objects.get(id=request.POST['subject_id'])

    # 大标题
    main_title = doc.add_paragraph(subject.name + '试卷')
    # 居中
    main_title.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    # 大标题样式
    main_title_style = doc.styles.add_style('MainTitleStyle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    main_title_style.font.size = docx.shared.Pt(20)
    main_title_style.font.name = u'宋体'
    main_title_style._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')
    main_title.style = main_title_style

    nameplace = doc.add_paragraph('学院_________专业_________班级_________姓名_________学号_________')
    nameplace.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    # 获取题目信息
    [choice_list, fill_list, judge_list, discuss_list] = json.loads(request.POST['questions'])

    index_str = ['一、', '二、', '三、', '四、']
    index_i = 0

    # 标题样式
    title_style = doc.styles.add_style('TitleStyle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = docx.shared.Pt(15)
    title_style.font.name = u'宋体'
    title_style._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')

    # 预存答案
    choice_answer, fill_answer, judge_answer, discuss_answer = [], [], [], []

    # 选择
    choice_num = len(choice_list)
    if choice_num > 0:
        # 空行
        doc.add_paragraph()
        # 标题
        doc.add_paragraph(index_str[index_i] + '选择题', style=title_style)

        index_i += 1

        # 选择题答题区域
        answer_indexs = []
        for i in range(choice_num // 5):
            begin = i * 5 + 1
            end = (i + 1) * 5
            answer_indexs.append(str(begin) + '-' + str(end) + ':')

        if choice_num % 5 != 0:
            answer_indexs.append(str(choice_num - choice_num % 5 + 1) + '-' + str(choice_num) + ':')

        for i in range(0, len(answer_indexs) - 1, 2):
            doc.add_paragraph(answer_indexs[i] + ' ' * 30 + answer_indexs[i + 1])

        if len(answer_indexs) % 2 == 1:
            doc.add_paragraph(answer_indexs[len(answer_indexs) - 1])

        # 选择题题目
        for i in range(choice_num):
            choice = Choice.objects.get(id=choice_list[i])
            doc.add_paragraph(str(i + 1) + '. ' + str.strip(choice.question))
            if len(choice.option_A + choice.option_B + choice.option_C + choice.option_D) < 25:
                doc.add_paragraph('A. ' + str.strip(choice.option_A) +
                                  '     B. ' + str.strip(choice.option_B) +
                                  '     C. ' + str.strip(choice.option_C) +
                                  '     D. ' + str.strip(choice.option_D))
            elif len(choice.option_A) < 15:
                doc.add_paragraph('A. ' + str.strip(choice.option_A) +
                                  '     B. ' + str.strip(choice.option_B))
                doc.add_paragraph('C. ' + str.strip(choice.option_C) +
                                  '     D. ' + str.strip(choice.option_D))
            else:
                doc.add_paragraph('A. ' + str.strip(choice.option_A))
                doc.add_paragraph('B. ' + str.strip(choice.option_B))
                doc.add_paragraph('C. ' + str.strip(choice.option_C))
                doc.add_paragraph('D. ' + str.strip(choice.option_D))

            choice_answer.append(choice.answer)

    # 填空
    if len(fill_list) > 0:
        doc.add_paragraph()
        doc.add_paragraph(index_str[index_i] + '填空题', style=title_style)

        index_i += 1

        for i in range(len(fill_list)):
            fill = Fill.objects.get(id=fill_list[i])
            doc.add_paragraph(str(i + 1) + '. ' + str.strip(fill.question))

            fill_answer.append(str.strip(fill.answer))

    # 判断
    if len(judge_list) > 0:
        doc.add_paragraph()
        doc.add_paragraph(index_str[index_i] + '判断题', style=title_style)

        index_i += 1

        for i in range(len(judge_list)):
            judge = Judge.objects.get(id=judge_list[i])
            doc.add_paragraph('(     )  ' + str(i + 1) + '. ' + str.strip(judge.question))

            judge_answer.append(judge.answer)

    # 简答
    if len(discuss_list) > 0:
        doc.add_paragraph()
        doc.add_paragraph(index_str[index_i] + '简答题', style=title_style)

        index_i += 1

        for i in range(len(discuss_list)):
            discuss = Discuss.objects.get(id=discuss_list[i])
            doc.add_paragraph(str(i + 1) + '. ' + str.strip(discuss.question))
            doc.add_paragraph()
            doc.add_paragraph()
            doc.add_paragraph()
            doc.add_paragraph()
            doc.add_paragraph()

            discuss_answer.append(str.strip(discuss.answer))

    # 后附答案
    doc.add_page_break()
    doc.add_paragraph('参考答案', style=title_style)

    doc.add_paragraph('选择题', style=title_style)
    para = doc.add_paragraph()
    for i in range(len(choice_answer)):
        para.add_run(str(i + 1) + '.' + choice_answer[i] + '   ')

    doc.add_paragraph('填空题', style=title_style)
    for i in range(len(fill_answer)):
        buffer = str(i + 1) + '.'
        answer_arr = json.loads(fill_answer[i])
        for ans in answer_arr:
            buffer += ' ' + ans

        doc.add_paragraph(buffer + '   ')

    doc.add_paragraph('判断题', style=title_style)
    para = doc.add_paragraph()
    for i in range(len(judge_answer)):
        para.add_run(str(i + 1) + '.' + judge_answer[i] + '   ')

    doc.add_paragraph('简答题', style=title_style)
    for i in range(len(discuss_answer)):
        doc.add_paragraph(str(i + 1) + '.' + discuss_answer[i])
        doc.add_paragraph()

    filedir = os.path.join(PAPER_DIR, user.username)
    if not os.path.exists(filedir):
        os.makedirs(filedir)
    filename = time.strftime('%Y%m%d-%H%M%S', time.localtime()) + '.docx'
    doc.save(os.path.join(filedir, filename))

    response = {'status': 'success', 'filename': filename}
    return JsonResponse(response)


def download(request):
    user = User.objects.get(username=request.GET['openid'])
    profile = UserProfile.objects.get(user=user)

    # 检测是否教师
    if not profile.isTeacher:
        response = {'status': 'fail', 'errMsg': 'permission denied.'}
        return JsonResponse(response)

    filedir = os.path.join(PAPER_DIR, user.username)
    filename = request.GET['name']
    paper = open(os.path.join(filedir, filename), 'rb')

    response = HttpResponse(paper.read(), content_type=DOCX_CONTENT_TYPE)
    response['Content-Disposition'] = 'attachment;filename="%s"' % filename
    return response


def get_list(request):
    user = User.objects.get(username=request.GET['openid'])
    profile = UserProfile.objects.get(user=user)

    # 检测是否教师
    if not profile.isTeacher:
        response = {'status': 'fail', 'errMsg': 'permission denied.'}
        return JsonResponse(response)

    filedir = os.path.join(PAPER_DIR, user.username)
    files = tuple(os.walk(filedir))[0][2]

    response = {'status': 'success', 'paper_list': files}
    return JsonResponse(response)
